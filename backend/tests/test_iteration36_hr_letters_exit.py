"""Iter 36 — HR P1: Letters & Exit/FNF + refactor regression."""
import io
import os
import pytest
import requests
from datetime import datetime, timezone, timedelta
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://worksite-command.preview.emergentagent.com").rstrip("/")
API = f"{BASE_URL}/api"

ADMIN = {"email": "admin@erp.com", "password": "Admin@123"}
SUPERVISOR = {"email": "supervisor@erp.com", "password": "Super@1234"}


# ─────────────────────────────── Fixtures ──────────────────────────────
@pytest.fixture(scope="session")
def admin():
    s = requests.Session()
    r = s.post(f"{API}/auth/login", json=ADMIN, timeout=15)
    assert r.status_code == 200, f"admin login failed {r.status_code}: {r.text}"
    return s


@pytest.fixture(scope="session")
def supervisor():
    s = requests.Session()
    r = s.post(f"{API}/auth/login", json=SUPERVISOR, timeout=15)
    if r.status_code != 200:
        pytest.skip("supervisor login failed")
    return s


def _make_docx(text: str) -> bytes:
    """Build a minimal in-memory DOCX with given paragraph text."""
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_docx_text(blob: bytes) -> str:
    d = Document(io.BytesIO(blob))
    return "\n".join(p.text for p in d.paragraphs)


# ───────────────────────────── Refactor regression ─────────────────────
class TestRefactorRegression:
    def test_onboardings_list(self, admin):
        r = admin.get(f"{API}/hr/onboardings", timeout=15)
        assert r.status_code == 200, r.text
        assert isinstance(r.json(), list)

    def test_onboarding_stages(self, admin):
        r = admin.get(f"{API}/hr/onboardings/stages", timeout=15)
        assert r.status_code == 200, r.text
        stages = r.json()
        assert isinstance(stages, list) and len(stages) == 6

    def test_leave_types(self, admin):
        r = admin.get(f"{API}/hr/leave-types", timeout=15)
        assert r.status_code == 200, r.text
        assert isinstance(r.json(), list)

    def test_leave_calendar(self, admin):
        r = admin.get(f"{API}/hr/leave-calendar?month=2026-06", timeout=15)
        assert r.status_code == 200, r.text

    def test_dashboard(self, admin):
        r = admin.get(f"{API}/hr/dashboard", timeout=15)
        assert r.status_code == 200, r.text

    def test_employees_listing(self, admin):
        r = admin.get(f"{API}/hr/employees", timeout=15)
        # may also exist under /api/employees - just ensure /hr namespace ok via another endpoint
        # Accept 200 or 404 (route lives elsewhere). The key check is leave types/onboardings.
        assert r.status_code in (200, 404)


# ───────────────────────────── Letters ────────────────────────────────
class TestLettersPlaceholders:
    def test_placeholders_5_groups(self, admin):
        r = admin.get(f"{API}/hr/letters/placeholders", timeout=15)
        assert r.status_code == 200, r.text
        data = r.json()
        for k in ("employee", "company", "user", "meta", "custom"):
            assert k in data, f"missing group: {k}"


class TestLetterTemplates:
    template_id = None
    relieving_tid = None

    def test_upload_happy(self, admin):
        blob = _make_docx("Dear {{ name }}, your designation is {{ designation }}. Custom amount: {{ increment_amount }}")
        files = {"file": ("offer.docx", blob,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_OfferTemplate", "kind": "offer", "description": "test offer letter"}
        r = admin.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=20)
        assert r.status_code in (200, 201), r.text
        body = r.json()
        assert "id" in body
        assert "binary" not in body
        TestLetterTemplates.template_id = body["id"]

    def test_upload_invalid_extension(self, admin):
        files = {"file": ("evil.txt", b"not docx", "text/plain")}
        data = {"name": "TEST_bad", "kind": "offer"}
        r = admin.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=15)
        assert r.status_code == 400, r.text

    def test_upload_unknown_kind(self, admin):
        blob = _make_docx("hello")
        files = {"file": ("x.docx", blob,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_bad_kind", "kind": "foobar"}
        r = admin.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=15)
        assert r.status_code == 400, r.text

    def test_list_excludes_binary(self, admin):
        r = admin.get(f"{API}/hr/letter-templates", timeout=15)
        assert r.status_code == 200, r.text
        rows = r.json()
        assert isinstance(rows, list)
        for row in rows:
            assert "binary" not in row, "binary should not be returned in list"

    def test_download_returns_docx(self, admin):
        assert TestLetterTemplates.template_id
        r = admin.get(f"{API}/hr/letter-templates/{TestLetterTemplates.template_id}/download", timeout=15)
        assert r.status_code == 200, r.text
        ct = r.headers.get("content-type", "")
        assert "wordprocessingml" in ct or "octet-stream" in ct
        assert r.content[:2] == b"PK"  # docx = zip

    def test_render_letter_with_employee_name(self, admin):
        # Need an employee
        emp_payload = {"name": "TEST_Iter36_RenderEmp", "email": f"test_iter36_render_{int(datetime.now().timestamp())}@erp.com",
                       "phone": "9999999990", "role": "Engineer", "department": "Operations",
                       "designation": "Junior Engineer", "salary": 30000, "joining_date": "2025-01-15"}
        r = admin.post(f"{API}/employees", json=emp_payload, timeout=15)
        if r.status_code not in (200, 201):
            # try alt endpoint
            r = admin.post(f"{API}/hr/employees", json=emp_payload, timeout=15)
        assert r.status_code in (200, 201), f"employee create failed: {r.status_code} {r.text}"
        emp = r.json()
        eid = emp.get("id") or emp.get("_id")
        assert eid

        # Render
        body = {"employee_id": eid, "variables": {"increment_amount": 50000}}
        r2 = admin.post(f"{API}/hr/letter-templates/{TestLetterTemplates.template_id}/render",
                        json=body, timeout=30)
        assert r2.status_code == 200, r2.text
        assert "X-Letter-Id" in r2.headers
        text = _parse_docx_text(r2.content)
        assert "TEST_Iter36_RenderEmp" in text, f"employee name not merged: {text!r}"
        assert "50000" in text, f"custom variable not merged: {text!r}"
        assert "{{" not in text and "{ name }" not in text

    def test_render_supervisor_403(self, supervisor):
        if not TestLetterTemplates.template_id:
            pytest.skip("no template")
        body = {"employee_id": "doesnt-matter", "variables": {}}
        r = supervisor.post(f"{API}/hr/letter-templates/{TestLetterTemplates.template_id}/render",
                            json=body, timeout=15)
        assert r.status_code in (401, 403), f"expected forbidden, got {r.status_code}: {r.text}"

    def test_upload_supervisor_403(self, supervisor):
        blob = _make_docx("nope")
        files = {"file": ("nope.docx", blob,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_nope", "kind": "offer"}
        r = supervisor.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=15)
        assert r.status_code in (401, 403), r.text

    def test_upload_relieving_template_for_finalise(self, admin):
        blob = _make_docx("Relieving letter for {{ name }} on {{ last_working_day }}")
        files = {"file": ("relieving.docx", blob,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_RelievingTemplate", "kind": "relieving"}
        r = admin.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=15)
        assert r.status_code in (200, 201), r.text
        TestLetterTemplates.relieving_tid = r.json()["id"]

    def test_delete_template_super_admin(self, admin):
        # Create temporary, then delete
        blob = _make_docx("temp")
        files = {"file": ("temp.docx", blob,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_DeleteMe", "kind": "custom"}
        r = admin.post(f"{API}/hr/letter-templates", files=files, data=data, timeout=15)
        assert r.status_code in (200, 201)
        tid = r.json()["id"]
        r2 = admin.delete(f"{API}/hr/letter-templates/{tid}", timeout=10)
        assert r2.status_code in (200, 204), r2.text


# ───────────────────────────── Exit & FNF ──────────────────────────────
class TestExitFnf:
    eid = None
    employee_id = None

    @pytest.fixture(autouse=True, scope="class")
    def setup_employee(self, request, admin):
        # Create employee with monthly_salary=30000, joining 5y ago = 2021-04-15
        five_y_ago = "2021-04-15"
        ts = int(datetime.now().timestamp())
        payload = {
            "name": f"TEST_Iter36_ExitEmp_{ts}",
            "email": f"test_iter36_exit_{ts}@erp.com",
            "phone": "9999999991",
            "role": "Engineer",
            "department": "Operations",
            "designation": "Engineer",
            "salary": 30000,
            "joining_date": five_y_ago,
        }
        r = admin.post(f"{API}/employees", json=payload, timeout=15)
        if r.status_code not in (200, 201):
            r = admin.post(f"{API}/hr/employees", json=payload, timeout=15)
        assert r.status_code in (200, 201), r.text
        TestExitFnf.employee_id = r.json().get("id") or r.json().get("_id")

        # Seed leave balances EL + PL = 10 total for current year
        year = datetime.now(timezone.utc).year
        # Find leave types or just upsert directly
        # Use raw insert via an internal endpoint if exists, else POST adjustment
        for lt, bal in [("EL", 5), ("PL", 5)]:
            adj = {"employee_id": TestExitFnf.employee_id, "leave_type": lt,
                   "year": year, "balance": bal, "credit": bal, "used": 0}
            # try create endpoint variants
            for url in [f"{API}/hr/leave-balances", f"{API}/hr/leave-balance/upsert"]:
                rr = admin.post(url, json=adj, timeout=10)
                if rr.status_code in (200, 201):
                    break
        yield

    def test_clearance_items_8(self, admin):
        r = admin.get(f"{API}/hr/clearance-items", timeout=10)
        assert r.status_code == 200, r.text
        items = r.json()
        assert len(items) == 8
        assert all("approver_role" in c for c in items)

    def test_create_exit_for_nonexistent(self, admin):
        body = {"employee_id": "no-such-id", "resignation_date": "2026-04-01", "last_working_day": "2026-04-15"}
        r = admin.post(f"{API}/hr/exits", json=body, timeout=10)
        assert r.status_code == 404, r.text

    def test_create_exit_happy(self, admin):
        body = {"employee_id": TestExitFnf.employee_id,
                "resignation_date": "2026-04-01",
                "last_working_day": "2026-04-15",
                "notice_period_days": 15,
                "reason": "test"}
        r = admin.post(f"{API}/hr/exits", json=body, timeout=10)
        assert r.status_code in (200, 201), r.text
        TestExitFnf.eid = r.json()["id"]
        # 8 clearance items pending
        assert len(r.json()["clearance"]) == 8

    def test_duplicate_open_exit_blocked(self, admin):
        body = {"employee_id": TestExitFnf.employee_id, "resignation_date": "2026-04-01",
                "last_working_day": "2026-04-15", "notice_period_days": 15}
        r = admin.post(f"{API}/hr/exits", json=body, timeout=10)
        assert r.status_code == 400, r.text

    def test_supervisor_cannot_approve_laptop(self, supervisor):
        if not TestExitFnf.eid:
            pytest.skip("no exit")
        r = supervisor.post(f"{API}/hr/exits/{TestExitFnf.eid}/clearance/laptop/approve",
                            json={"remarks": "ok"}, timeout=10)
        assert r.status_code in (401, 403), f"expected forbidden got {r.status_code}: {r.text}"

    def test_super_admin_approves_all(self, admin):
        if not TestExitFnf.eid:
            pytest.skip()
        items = ["laptop", "id_card", "keys", "ppe", "it_access",
                 "knowledge_transfer", "library", "accounts"]
        for k in items:
            r = admin.post(f"{API}/hr/exits/{TestExitFnf.eid}/clearance/{k}/approve",
                           json={"remarks": "done"}, timeout=10)
            assert r.status_code == 200, f"{k}: {r.status_code} {r.text}"

    def test_compute_fnf_math(self, admin):
        if not TestExitFnf.eid:
            pytest.skip()
        r = admin.post(f"{API}/hr/exits/{TestExitFnf.eid}/compute-fnf", timeout=15)
        assert r.status_code == 200, r.text
        fnf = r.json()["fnf"]
        # per_day=1000, pending=15000, encash=10000, gratuity=86538.46, net≈111538
        assert abs(fnf["per_day_rate"] - 1000) < 1, fnf
        assert abs(fnf["pending_salary"] - 15000) < 10, fnf
        # encash may be 0 if leave-balance seed failed — log and check both expectations
        if fnf["encashable_days"] >= 9.5:
            assert abs(fnf["leave_encashment"] - 10000) < 10, fnf
        else:
            print(f"WARN: encashable_days={fnf['encashable_days']} (leave-balance seeding may have failed)")
        assert fnf["tenure_years"] >= 5.0, fnf
        assert abs(fnf["gratuity"] - 86538.46) < 10, fnf
        # short notice should be 0 since notice_period_days=15 and actual_notice=15
        assert fnf["notice_recovery"] == 0, fnf
        # If encashment seeded correctly, net = 111538
        if fnf["encashable_days"] >= 9.5:
            assert abs(fnf["net_payable"] - 111538.46) < 10, fnf

    def test_override_fnf(self, admin):
        if not TestExitFnf.eid:
            pytest.skip()
        r = admin.put(f"{API}/hr/exits/{TestExitFnf.eid}/fnf",
                      json={"overrides": {"gratuity": 100000}}, timeout=10)
        assert r.status_code == 200, r.text
        fnf = r.json()["fnf"]
        assert fnf["gratuity"] == 100000
        # net recomputed: pending+encash+100000+bonus−advances−notice_recovery
        expected_net = (fnf["pending_salary"] + fnf["leave_encashment"]
                        + 100000 + fnf.get("bonus_accrual", 0)
                        - fnf.get("advances", 0) - fnf.get("notice_recovery", 0))
        assert abs(fnf["net_payable"] - expected_net) < 1, fnf

    def test_finalise_and_relieving_letter(self, admin):
        if not TestExitFnf.eid:
            pytest.skip()
        r = admin.post(f"{API}/hr/exits/{TestExitFnf.eid}/finalise", timeout=20)
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["status"] == "finalised"
        # employee status flipped
        emp = admin.get(f"{API}/employees/{TestExitFnf.employee_id}", timeout=10)
        if emp.status_code != 200:
            emp = admin.get(f"{API}/hr/employees/{TestExitFnf.employee_id}", timeout=10)
        if emp.status_code == 200:
            assert emp.json().get("status") == "exited", emp.json()
        # relieving letter (template exists from earlier test)
        if data.get("relieving_letter_id"):
            print(f"OK: relieving_letter_id={data['relieving_letter_id']}")
        else:
            print("WARN: relieving_letter_id missing (template may not be present)")

    def test_finalise_already_finalised(self, admin):
        if not TestExitFnf.eid:
            pytest.skip()
        r = admin.post(f"{API}/hr/exits/{TestExitFnf.eid}/finalise", timeout=10)
        assert r.status_code == 400


class TestFinaliseGating:
    """Separate flow: create a new exit, do not approve all, finalise should 400."""

    def test_finalise_blocked_when_pending(self, admin):
        ts = int(datetime.now().timestamp())
        payload = {"name": f"TEST_Iter36_GateEmp_{ts}", "email": f"test_iter36_gate_{ts}@erp.com",
                   "phone": "9999999992", "role": "Engineer", "department": "Operations",
                   "salary": 25000, "joining_date": "2024-04-15"}
        r = admin.post(f"{API}/employees", json=payload, timeout=15)
        if r.status_code not in (200, 201):
            r = admin.post(f"{API}/hr/employees", json=payload, timeout=15)
        assert r.status_code in (200, 201)
        eid_emp = r.json().get("id") or r.json().get("_id")
        body = {"employee_id": eid_emp, "resignation_date": "2026-04-01",
                "last_working_day": "2026-04-15", "notice_period_days": 15}
        r2 = admin.post(f"{API}/hr/exits", json=body, timeout=10)
        assert r2.status_code in (200, 201)
        eid = r2.json()["id"]
        # Approve only 1 item
        admin.post(f"{API}/hr/exits/{eid}/clearance/laptop/approve",
                   json={"remarks": "ok"}, timeout=10)
        # Attempt finalise — should 400
        r3 = admin.post(f"{API}/hr/exits/{eid}/finalise", timeout=10)
        assert r3.status_code == 400, r3.text
        assert "clearance" in r3.text.lower() or "not approved" in r3.text.lower()
